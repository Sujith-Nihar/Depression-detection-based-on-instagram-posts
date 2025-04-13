import pandas as pd
import google.generativeai as genai
from dotenv import load_dotenv
import os
from docx import Document
import base64
import mimetypes

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

# Load the dataset
merged_df = pd.read_excel('/Users/sujiththota/Downloads/Python/Research/southernconstellation_results/southernconstellation_Mood_filtered.xlsx')

# Function to extract timestamp from the filename
def extract_timestamp(filename):
    try:
        date_time_part = filename.split("_UTC")[0]  # Remove "_UTC_x.jpg"
        date, time = date_time_part.split("_")  # Split into date and time
        time = time.replace("-", ":")  # Replace hyphens in time with colons
        return f"{date} {time}"
    except Exception as e:
        print(f"Error processing filename '{filename}': {e}")
        return None


def encode_video(video_path):
    """
    Encodes a video file as a base64 string for API input.
    """
    try:
        with open(video_path, "rb") as video_file:
            return base64.b64encode(video_file.read()).decode("utf-8")
    except FileNotFoundError:
        print(f"Warning: Video not found - {video_path}")
        return None

# Apply timestamp extraction
merged_df['Timestamp'] = merged_df['Media Name'].apply(extract_timestamp)

# Convert to datetime and sort by timestamp
merged_df['Timestamp'] = pd.to_datetime(merged_df['Timestamp'], format="%Y-%m-%d %H:%M:%S")
sorted_df = merged_df.sort_values(by='Timestamp')
genai.configure(api_key=api_key)

def encode_file(file_path):
    """
    Encodes an image or video file as a base64 string for API input.
    """
    try:
        with open(file_path, "rb") as file:
            return base64.b64encode(file.read()).decode("utf-8")
    except FileNotFoundError:
        print(f"Warning: File not found - {file_path}")
        return None
    


def summarize_mood_course_gemini(posts, media_folder="/Users/sujiththota/Downloads/Python/Research/southernconstellation_phq/"):
    if posts.empty:
        return "No significant emotional trajectory detected."

    model = genai.GenerativeModel("gemini-1.5-flash")

    # Prompt for analysis
    prompt = """
    You are an AI psychiatrist analyzing a user's Instagram posts over time. 
    Each post contains an Timestamp, image/video, embedded text, caption.
    
    **Your Task:**
    Analyze the image/video, caption, embedded text with timestamp and do the following
    1) Summarize the user's mood course and temporal pattern over time (mood is cetogorized as Happiness, Sadness, Fear, Disgust, Anger, Surprise)
    2) progression of emotional states and create a trajectory of progression of the mental illness if present over time
    3) Overal Analysis based on all features

    (summarize in time periods, choose time periods by events present) consider all events but also stress topics like health related issues, diagnosis, events.:
    
    **Data for Analysis:**
    """

    # Start with the prompt
    content_parts = [{"text": prompt}]

    for _, row in posts.iterrows():
        media_path = os.path.join(media_folder, row['Media Name'])
        media_type, _ = mimetypes.guess_type(media_path)

        # Base text for each row
        row_text = f"""
        **Timestamp:** {row['Timestamp']}
        - **Media Type:** {"Video" if media_type and "video" in media_type else "Image"}
        - **Description:** {row['Simple Description']}
        - **Embedded Text:** {row['Embedded Text']}
        - **Caption:** {row['Media Caption']}
        """

        # Append text data
        content_parts.append({"text": row_text})

        # Handle images (encode and attach)
        if media_type and "image" in media_type:
            image_data = encode_file(media_path)
            if image_data:
                content_parts.append({
                    "inline_data": {
                        "mime_type": media_type,
                        "data": image_data
                    }
                })

        # Handle videos (encode and attach)
        elif media_type and "video" in media_type:
            video_data = encode_video(media_path)
            if video_data:
                content_parts.append({
                        "mime_type": "video/mp4",
                        "data": video_data
                })


    # Generate response using Gemini
    response = model.generate_content(content_parts, generation_config=genai.types.GenerationConfig(temperature=0.1))

    return response.text

# Function to save the summary as a Word document
def save_summary_to_docx(summary, filename='/Users/sujiththota/Downloads/Python/Research/southernconstellation_results/southernconstellation_mood_summary.docx'):
    doc = Document()
    doc.add_heading('Mood Summary', level=1)
    doc.add_paragraph(summary)
    doc.save(filename)
    print(f"Mood summary saved to '{filename}'.")

# Run the summarization
mood_summary = summarize_mood_course_gemini(sorted_df)

# Save the summary
save_summary_to_docx(mood_summary)
