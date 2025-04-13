import pandas as pd
import google.generativeai as genai
from dotenv import load_dotenv
import os
from docx import Document
import base64
import mimetypes

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

# Load the dataset
merged_df = pd.read_excel('/Users/sujiththota/Downloads/Python/Research/southernconstellation_results/southernconstellation_PHQ_filtered.xlsx')

# Function to extract timestamp from the filename
def extract_timestamp(filename):
    """
    Extracts the timestamp as a string from a filename in the format: YYYY-MM-DD_HH-MM-SS_UTC_x.jpg.
    """
    try:
        date_time_part = filename.split("_UTC")[0]  # Remove "_UTC_x.jpg"
        date, time = date_time_part.split("_")  # Split into date and time
        time = time.replace("-", ":")  # Replace hyphens in time with colons
        return f"{date} {time}"
    except Exception as e:
        print(f"Error processing filename '{filename}': {e}")
        return None

# Apply timestamp extraction
merged_df['Timestamp'] = merged_df['Media Name'].apply(extract_timestamp)

# Convert to datetime and sort by timestamp
merged_df['Timestamp'] = pd.to_datetime(merged_df['Timestamp'], format="%Y-%m-%d %H:%M:%S")
sorted_df = merged_df.sort_values(by='Timestamp')

# Save the sorted dataset
# sorted_df.to_excel('/Users/sujiththota/Downloads/Python/Research/catherineanne0201_menopause_results/catherineanne0201_menopause_PHQ_9_with_timestamp.xlsx', index=False)

# Configure Gemini API
genai.configure(api_key=api_key)


def encode_image(image_path):
    """
    Encodes an image file as a base64 string for API input.
    """
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode("utf-8")
    except FileNotFoundError:
        print(f"Warning: Image not found - {image_path}")
        return None
    

def encode_video(video_path):
    """
    Encodes a video file as a base64 string for API input.
    """
    try:
        with open(video_path, "rb") as video_file:
            return base64.b64encode(video_file.read()).decode("utf-8")
    except FileNotFoundError:
        print(f"Warning: Video not found - {video_path}")
        return None

# Function to analyze PHQ-9 progression
def summarize_phq_course_gemini(posts, media_folder="/Users/sujiththota/Downloads/Python/Research/southernconstellation_phq/"):
    if posts.empty:
        return "No significant emotional trajectory detected."

    model = genai.GenerativeModel("gemini-1.5-flash")

    # Prompt for analysis
    prompt = """ 
    You are an AI psychiatrist analyzing a user's Instagram posts over time. 
    Each post contains an Timestamp, image/video, embedded text, caption, and PHQ-9 symptoms.
    
    **Your Task:**
    Analyze the image/video, caption, embedded text and PHQ-9 symptoms with timestamp and do the following:  
    1) summmarize the user's PHQ-9 progression and temporal pattern over time (summarize in time periods, choose time periods by events present)
    2) create a trajectory of progression of the mental illness over time in paragraph not table(I want a section progression of PHQ-9)
    based on all the features I have provided over time: and add justification from which feature you drawn the conclusion (for example phq-9 or image or caption)
    3) Overal Analysis based on all features
    
    **Data for Analysis:**
    """


    content_parts = [{"text": prompt}]

    for _, row in posts.iterrows():
        media_path = os.path.join(media_folder, row['Media Name'])
        media_type, _ = mimetypes.guess_type(media_path)


        row_text = f"""
        **Timestamp:** {row['Timestamp']}
        - **Image/Video Description:** {row['Simple Description']}
        - **Embedded Text:** {row['Embedded Text']}
        - **Caption:** {row['Media Caption']}
        - **PHQ-9 Symptoms:**
          - Loss of Interest: {row['Loss of Interest Binary']}
          - Feeling Depressed: {row['Feeling depressed Binary']}
          - Sleeping Disorder: {row['Sleeping Disorder Binary']}
          - Lack of Energy: {row['Lack of Energy Binary']}
          - Eating Disorder: {row['Eating Disorder Binary']}
          - Low Self-Esteem: {row['Low Self-Esteem Binary']}
          - Concentration Difficulty: {row['Concentration difficulty Binary']}
          - Psychomotor Changes: {row['Psychomotor changes Binary']}
          - Self-Harm Risk: {row['Self harm risk Binary']}
        """


        content_parts.append({"text": row_text})


        # Handling images (encode and attach)
        if media_type and "image" in media_type:
            image_data = encode_image(media_path)
            if image_data:
                content_parts.append({
                    "inline_data": {
                        "mime_type": media_type,
                        "data": image_data
                    }
                })

         # Handling videos (encode and attach)
        elif media_type and "video" in media_type:
            video_data = encode_video(media_path)
            if video_data:
                content_parts.append({
                        "mime_type": "video/mp4",
                        "data": video_data
                })

    response = model.generate_content(content_parts, generation_config=genai.types.GenerationConfig(temperature=0.1))

    return response.text

# Function to save the summary as a Word document
def save_summary_to_docx(summary, filename='/Users/sujiththota/Downloads/Python/Research/southernconstellation_results/southernconstellation_PHQ_9_SUMMARY_test.docx'):
    doc = Document()
    doc.add_heading('PHQ-9 Summary', level=1)
    doc.add_paragraph(summary)
    doc.save(filename)
    print(f"PHQ-9 summary saved to '{filename}'.")

# Run the summarization
mood_summary = summarize_phq_course_gemini(sorted_df)

# Save the summary
save_summary_to_docx(mood_summary)
